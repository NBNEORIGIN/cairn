"""Voice surface — file upload + text extraction.

Receives one or more files via multipart/form-data and returns the
extracted text per file. The Next.js /voice page uploads files
attached to a chat message; the extracted text is then prepended to
the user's prompt before forwarding to the agent stream.

No file is persisted by default — extraction is in-memory and the
blob is discarded once the response is sent. If the request includes
``persist=true``, the original bytes are written to the volume at
``/app/data/uploads/<sha8>/<filename>`` for later retrieval (defer
this hook until we actually need it).

Supported formats:
    .pdf            — pypdf
    .docx           — python-docx
    .csv / .tsv     — stdlib csv (capped at 1k rows + summary)
    .xlsx / .xlsm   — openpyxl (first sheet, capped at 1k rows)
    .txt / .md      — plain decode
    .png/.jpg/.gif/.webp — base64-encoded for the chat path's
                          image_base64 field (Claude/OpenAI vision)
    .heic / .bmp    — accepted, converted to JPEG via Pillow before
                      base64 encoding (Anthropic's vision API doesn't
                      accept HEIC/BMP natively)

Per-file size cap: 10 MB. Per-request: 5 files. Hard caps so the
agent's context isn't drowned by a stray 50MB report.
"""
from __future__ import annotations

import base64
import csv
import hashlib
import io
import logging
from typing import Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse

from api.middleware.auth import verify_api_key

log = logging.getLogger(__name__)
router = APIRouter(prefix='/voice', tags=['Voice file upload'])


MAX_FILE_BYTES = 10 * 1024 * 1024     # 10 MB
MAX_FILES_PER_REQUEST = 5
MAX_TEXT_CHARS = 60_000               # ~15k tokens per file
MAX_CSV_ROWS = 1_000
MAX_XLSX_ROWS = 1_000


def _sha8(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:8]


def _ext(filename: str) -> str:
    return ('.' + (filename or '').rsplit('.', 1)[-1].lower()) if '.' in (filename or '') else ''


def _truncate(s: str, limit: int = MAX_TEXT_CHARS) -> tuple[str, bool]:
    if len(s) <= limit:
        return s, False
    return s[:limit] + '\n…[truncated]', True


def _extract_pdf(blob: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob))
    out: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ''
        except Exception as exc:
            text = f'[page {i + 1} unreadable: {type(exc).__name__}]'
        out.append(f'--- page {i + 1} ---\n{text.strip()}')
    return '\n\n'.join(out).strip()


def _extract_docx(blob: bytes) -> str:
    import docx as docx_mod  # python-docx
    doc = docx_mod.Document(io.BytesIO(blob))
    parts: list[str] = []
    for p in doc.paragraphs:
        line = (p.text or '').strip()
        if line:
            parts.append(line)
    # Tables
    for ti, table in enumerate(doc.tables):
        parts.append(f'\n[table {ti + 1}]')
        for row in table.rows:
            cells = [(c.text or '').strip() for c in row.cells]
            parts.append(' | '.join(cells))
    return '\n'.join(parts).strip()


def _extract_csv(blob: bytes, sep: str = ',') -> str:
    text = blob.decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(text), delimiter=sep)
    rows: list[list[str]] = []
    for i, row in enumerate(reader):
        if i >= MAX_CSV_ROWS:
            rows.append(['…', f'(truncated at {MAX_CSV_ROWS} rows)'])
            break
        rows.append(row)
    if not rows:
        return '[empty CSV]'
    header = rows[0]
    body = rows[1:]
    summary = (
        f'columns: {len(header)} | rows: {len(body)} | '
        f'header: {", ".join(header[:10])}{"…" if len(header) > 10 else ""}'
    )
    sample_lines = [' | '.join(r) for r in rows[:50]]
    return summary + '\n\n' + '\n'.join(sample_lines)


def _extract_xlsx(blob: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(filename=io.BytesIO(blob), read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        return '[empty workbook]'
    rows: list[list[str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= MAX_XLSX_ROWS:
            rows.append(['…', f'(truncated at {MAX_XLSX_ROWS} rows)'])
            break
        rows.append([('' if c is None else str(c)) for c in row])
    wb.close()
    if not rows:
        return '[empty sheet]'
    header = rows[0]
    body = rows[1:]
    summary = (
        f'sheet: {ws.title} | columns: {len(header)} | rows: {len(body)} | '
        f'header: {", ".join(header[:10])}{"…" if len(header) > 10 else ""}'
    )
    sample_lines = [' | '.join(r) for r in rows[:50]]
    return summary + '\n\n' + '\n'.join(sample_lines)


def _extract_text(blob: bytes) -> str:
    return blob.decode('utf-8', errors='replace')


_EXTRACTORS: dict[str, Any] = {
    '.pdf': _extract_pdf,
    '.docx': _extract_docx,
    '.csv': lambda b: _extract_csv(b, ','),
    '.tsv': lambda b: _extract_csv(b, '\t'),
    '.xlsx': _extract_xlsx,
    '.xlsm': _extract_xlsx,
    '.txt': _extract_text,
    '.md': _extract_text,
    '.log': _extract_text,
    '.json': _extract_text,
}

_IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.heic', '.gif', '.webp', '.bmp'}

# Anthropic vision API media types — anything else needs conversion via Pillow.
# https://docs.anthropic.com/claude/docs/vision
_NATIVE_IMAGE_MEDIA_TYPES: dict[str, str] = {
    '.png':  'image/png',
    '.jpg':  'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.gif':  'image/gif',
    '.webp': 'image/webp',
}
_NEEDS_CONVERSION = {'.heic', '.bmp'}


def _encode_image(blob: bytes, ext: str) -> tuple[str, str]:
    """Return (base64_str, media_type). Converts HEIC/BMP to JPEG via Pillow.

    Raises on unsupported extensions or conversion failure — caller maps
    that to an error chip on the UI.
    """
    if ext in _NATIVE_IMAGE_MEDIA_TYPES:
        return base64.b64encode(blob).decode('ascii'), _NATIVE_IMAGE_MEDIA_TYPES[ext]
    if ext in _NEEDS_CONVERSION:
        # HEIC needs pillow-heif registered before PIL.Image.open() will
        # recognise it. BMP is handled by base Pillow but Anthropic's API
        # doesn't accept image/bmp — convert to JPEG either way.
        from PIL import Image
        if ext == '.heic':
            try:
                from pillow_heif import register_heif_opener
                register_heif_opener()
            except ImportError as exc:
                raise RuntimeError(
                    'HEIC support requires pillow-heif (pip install pillow-heif)'
                ) from exc
        img = Image.open(io.BytesIO(blob))
        # Anthropic JPEG: drop alpha, convert RGBA → RGB on white.
        if img.mode in ('RGBA', 'LA', 'P'):
            from PIL import Image as PILImage
            bg = PILImage.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        out_buf = io.BytesIO()
        img.save(out_buf, format='JPEG', quality=85, optimize=True)
        return base64.b64encode(out_buf.getvalue()).decode('ascii'), 'image/jpeg'
    raise ValueError(f'unsupported image extension: {ext}')


@router.post('/upload')
async def voice_upload(
    files: list[UploadFile] = File(...),
    _: bool = Depends(verify_api_key),
) -> JSONResponse:
    """Extract text from one or more uploaded files.

    Returns:
        {
          "files": [
            {
              "name": "report.pdf",
              "size": 123456,
              "type": ".pdf",
              "sha8": "a1b2c3d4",
              "text": "...extracted...",
              "truncated": false,
              "supported": true,
              "error": null
            },
            ...
          ]
        }
    """
    if not files:
        raise HTTPException(status_code=400, detail='no files')
    if len(files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(
            status_code=400,
            detail=f'too many files (max {MAX_FILES_PER_REQUEST})',
        )

    results: list[dict] = []
    for upload in files:
        name = upload.filename or 'upload'
        ext = _ext(name)
        size = 0
        text = ''
        truncated = False
        supported = ext in _EXTRACTORS or ext in _IMAGE_EXTS
        error: str | None = None
        # Vision payload — populated only for image attachments.
        image_base64: str | None = None
        image_media_type: str | None = None

        try:
            blob = await upload.read()
            size = len(blob)
            if size > MAX_FILE_BYTES:
                error = f'file exceeds {MAX_FILE_BYTES // (1024 * 1024)} MB limit'
                supported = False
            elif ext in _IMAGE_EXTS:
                try:
                    image_base64, image_media_type = _encode_image(blob, ext)
                    # Tell the chat path the model should look at the image.
                    text = f'[Image attachment: {name} ({image_media_type})]'
                except Exception as exc:
                    error = f'image encode failed: {type(exc).__name__}: {exc}'
                    supported = False
                    log.warning('[voice/upload] %s image encode failed: %s', name, exc)
            elif ext in _EXTRACTORS:
                extractor = _EXTRACTORS[ext]
                raw = extractor(blob)
                text, truncated = _truncate(raw or '')
            else:
                error = f'unsupported extension: {ext or "(none)"}'
                supported = False
        except Exception as exc:
            error = f'{type(exc).__name__}: {exc}'
            log.warning('[voice/upload] %s failed: %s', name, exc)
        finally:
            try:
                await upload.close()
            except Exception:
                pass

        results.append({
            'name': name,
            'size': size,
            'type': ext,
            'sha8': _sha8(blob) if size else '',
            'text': text,
            'truncated': truncated,
            'supported': supported,
            'error': error,
            # Vision: only set on image uploads, otherwise null. The
            # /voice page collects these and forwards them on the chat
            # payload so the agent's MessageEnvelope.image_base64 is
            # populated and Claude/OpenAI vision can see them.
            'image_base64': image_base64,
            'image_media_type': image_media_type,
        })

    return JSONResponse({'files': results})
